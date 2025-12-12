import sys
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_txt_to_docx(input_file, output_file):
    """
    Converts a plain text file to a .docx file, treating blank lines
    as paragraph separators. It also applies some basic styling.
    """
    if not os.path.isfile(input_file):
        print(f"Error: Input file not found at '{input_file}'")
        sys.exit(1)

    print(f"Reading input file: '{input_file}'...")
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        print(f"Error reading file: {e}")
        sys.exit(1)

    # Normalize newlines to ensure consistent paragraph separation.
    # Any sequence of 2 or more newlines becomes exactly two.
    normalized_content = re.sub(r'\n{2,}', '\n\n', content.strip())
    
    # Split the text into paragraphs based on the double newline.
    paragraphs_text = normalized_content.split('\n\n')

    # Create a new Word document
    document = Document()
    
    # Set basic styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    print(f"Creating .docx document from {len(paragraphs_text)} paragraph(s)...")
    
    for para_text in paragraphs_text:
        # Don't add empty strings as paragraphs.
        stripped_para = para_text.strip()
        if stripped_para:
            p = document.add_paragraph()
            # Add text
            p.add_run(stripped_para)
            # Justify alignment
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Set paragraph spacing (optional, but good practice)
            paragraph_format = p.paragraph_format
            paragraph_format.space_after = Pt(0) # No extra space after paragraph
            paragraph_format.line_spacing = 1.15

    # Save the document
    try:
        document.save(output_file)
        print(f"Successfully saved .docx file to: '{output_file}'")
    except Exception as e:
        print(f"Error saving .docx file: {e}")
        sys.exit(1)

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("\nUsage: python3 convert_to_docx.py <input_file.txt> <output_file.docx>\n")
        print("Example:")
        print("  python3 convert_to_docx.py text/chapters/prologue/ru.txt chapter_prologue.docx\n")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not output_path.lower().endswith('.docx'):
        print("Error: Output file must have a .docx extension.")
        sys.exit(1)

    convert_txt_to_docx(input_path, output_path)
