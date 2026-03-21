import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert_txt_to_docx(input_file: str | Path, output_file: str | Path) -> None:
    """
    Converts a plain text file to a .docx file, treating blank lines
    as paragraph separators. It also applies some basic styling.
    """
    input_file = Path(input_file)
    if not input_file.is_file():
        raise FileNotFoundError(f"Input file not found: '{input_file}'")

    try:
        content = input_file.read_text(encoding='utf-8')
    except OSError as e:
        raise OSError(f"Error reading file '{input_file}': {e}") from e

    # Normalize newlines: any sequence of 2+ newlines becomes exactly two
    normalized_content = re.sub(r'\n{2,}', '\n\n', content.strip())
    paragraphs_text = normalized_content.split('\n\n')

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for para_text in paragraphs_text:
        stripped_para = para_text.strip()
        if stripped_para:
            p = document.add_paragraph()
            p.add_run(stripped_para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph_format = p.paragraph_format
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = 1.15

    try:
        document.save(str(output_file))
    except OSError as e:
        raise OSError(f"Error saving .docx file to '{output_file}': {e}") from e
